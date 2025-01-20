from docx import Document
from datetime import datetime
from pathlib import Path
from fastapi import FastAPI, APIRouter, HTTPException
from fastapi.responses import FileResponse
import os

# Initialize the FastAPI app
app = FastAPI()
router = APIRouter()

# Define directories
BASE_DIR = Path("src/assets/permissions")
BASE_DIR.mkdir(parents=True, exist_ok=True)  # Ensure directory exists


# PDF Generation Function
def generate_permission_pdf(operation_type, items, location, receiver, delivery_person, manager_name):
    """
    Generate a permission PDF based on the given data.
    """
    template_path = Path("dn_template.docx")
    if not template_path.exists():
        raise FileNotFoundError("Template file not found!")

    # Load the Word template
    doc = Document(template_path)

    # Replace placeholders with actual data
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace("…………./……………/………………..", datetime.now().strftime("%Y-%m-%d"))
        paragraph.text = paragraph.text.replace("……………………………………………………………………………………………………………………………………", receiver)
        paragraph.text = paragraph.text.replace("…………………………………………………………………………………………………………………………………...", location)

    # Add items to the document
    items_table = doc.add_table(rows=1, cols=2)
    items_table.style = 'Table Grid'
    hdr_cells = items_table.rows[0].cells
    hdr_cells[0].text = 'Item Description'
    hdr_cells[1].text = 'Quantity'

    for item in items:
        row_cells = items_table.add_row().cells
        row_cells[0].text = item['description']
        row_cells[1].text = str(item['quantity'])

    # Save the document as a temporary Word file
    temp_doc_path = Path("temp.docx")
    doc.save(temp_doc_path)

    # Convert to PDF (placeholder for now, requires third-party library)
    pdf_path = BASE_DIR / f"{operation_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    convert_to_pdf(temp_doc_path, pdf_path)

    # Clean up temporary Word file
    temp_doc_path.unlink()

    return pdf_path


# Placeholder for DOCX to PDF conversion
def convert_to_pdf(docx_path, pdf_path):
    """
    Convert a DOCX file to a PDF file.
    """
    # Third-party conversion library needed here, e.g., docx2pdf or libreoffice.
    pass


# API Endpoint for PDF Generation
@router.post("/stock/generate_permission")
async def create_permission(
    operation_type: str,
    items: list[dict],
    location: str,
    receiver: str,
    delivery_person: str,
    manager_name: str
):
    """
    Generate a stock permission PDF and return the file path.
    """
    try:
        pdf_path = generate_permission_pdf(operation_type, items, location, receiver, delivery_person, manager_name)
        return {"file_path": str(pdf_path)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# API Endpoint for Downloading PDF
@router.get("/stock/permission/{filename}")
async def download_permission(filename: str):
    """
    Provide a download link for the generated PDF.
    """
    file_path = BASE_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path, filename=filename)

# Register the router
app.include_router(router)

# Ensure template exists for development purposes
if not Path("template.docx").exists():
    print("Warning: template.docx file is missing. Please add it to the project directory.")
